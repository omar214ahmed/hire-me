import io
from typing import Optional, Tuple


# =========================
# Public API
# =========================

def extract_pdf(file_bytes: bytes) -> Tuple[str, str, Optional[int]]:
    """
    Extract text from PDF bytes using:
    1. pdfplumber (primary)
    2. PyMuPDF (fallback)

    Returns:
        (text, parser_used, page_count)
    """

    text, page_count = _extract_with_pdfplumber(file_bytes)

    if _is_valid_text(text):
        return text, "pdfplumber", page_count

    text, page_count = _extract_with_pymupdf(file_bytes)

    if _is_valid_text(text):
        return text, "pymupdf", page_count

    return text or "", "empty_or_scanned_pdf", page_count


def extract_docx(file_bytes: bytes) -> Tuple[str, str, Optional[int]]:
    """
    Extract text from DOCX (paragraphs + tables)

    Returns:
        (text, parser_used, page_count)

    DOCX files don't store a reliable page count without rendering, so
    this is estimated from paragraph volume (~45 lines/page — good enough
    for a sanity-check cap, not for precise pagination).
    """

    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))

        parts = []

        # paragraphs
        parts.extend(_extract_docx_paragraphs(doc))

        # tables
        parts.extend(_extract_docx_tables(doc))

        text = "\n".join(parts)

        page_count = _estimate_docx_page_count(doc)

        return text, "python-docx", page_count

    except Exception as e:
        return "", f"docx_error:{type(e).__name__}", None


# =========================
# PDF implementations
# =========================

def _extract_with_pdfplumber(file_bytes: bytes) -> Tuple[Optional[str], Optional[int]]:
    try:
        import pdfplumber

        pages = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

        return "\n\n".join(pages), page_count

    except Exception:
        return None, None


def _extract_with_pymupdf(file_bytes: bytes) -> Tuple[Optional[str], Optional[int]]:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")

        page_count = doc.page_count
        pages = [page.get_text("text") for page in doc]

        doc.close()

        return "\n\n".join(pages), page_count

    except Exception:
        return None, None


# =========================
# DOCX helpers
# =========================

def _extract_docx_paragraphs(doc) -> list[str]:
    return [
        p.text.strip()
        for p in doc.paragraphs
        if p.text and p.text.strip()
    ]


def _extract_docx_tables(doc) -> list[str]:
    rows_text = []

    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            )

            if row_text:
                rows_text.append(row_text)

    return rows_text


def _estimate_docx_page_count(doc) -> int:
    """
    Rough page estimate: DOCX has no page-count metadata without a render
    pass, so approximate from paragraph line volume (~45 lines/page).
    """
    line_count = sum(1 for p in doc.paragraphs if p.text and p.text.strip())
    return max(1, -(-line_count // 45))  # ceil division


# =========================
# Utilities
# =========================

def _is_valid_text(text: Optional[str]) -> bool:
    return bool(text and len(text.strip()) > 50)