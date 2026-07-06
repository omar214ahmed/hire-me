"""
In-process integration test for the full HTTP surface the frontend
(ats_console.html) drives: create job -> upload candidate -> run match
-> poll task -> read shortlist/results.

Postgres, Redis, and the real ML models (GLiNER / BGE-M3 / reranker) are
mocked out so this runs anywhere with no external services — it's
checking that the routers, schemas, and JSON shapes are wired together
correctly end-to-end, not model quality.

Run with:  pytest tests/test_api_integration.py -v
"""
import io
import sys
import uuid
from pathlib import Path
from datetime import datetime, timezone

import pytest
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))


# ---------------------------------------------------------------------
# Fake in-memory storage (replaces pipeline/storage.py's Postgres calls)
# ---------------------------------------------------------------------
class FakeStorage:
    def __init__(self):
        self.jobs = {}
        self.candidates = {}

    async def save_job(self, jd_text, jd_extracted, jd_query, jd_embedding=None):
        job_id = uuid.uuid4().hex[:12]
        self.jobs[job_id] = {
            "job_id": job_id,
            "jd_text": jd_text,
            "extracted": jd_extracted,
            "query": jd_query,
            "jd_embedding": jd_embedding,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        return job_id

    async def get_job(self, job_id):
        return self.jobs.get(job_id)

    async def list_jobs(self):
        return list(self.jobs.values())

    async def save_candidate(self, file_id, filename, cv_parsed, cv_embedding=None):
        self.candidates[file_id] = {
            "filename": filename,
            "parsed": cv_parsed,
            "cv_embedding": cv_embedding,
        }
        return file_id

    async def get_candidate(self, candidate_id):
        return self.candidates.get(candidate_id)

    async def get_top_candidates_by_similarity(self, jd_embedding, limit):
        records = list(self.candidates.values())[:limit]
        for r in records:
            r["semantic_score"] = 0.5
        return records


fake_storage = FakeStorage()


# ---------------------------------------------------------------------
# Fake ML models (replaces pipeline/models.py::ModelRegistry)
# ---------------------------------------------------------------------
class FakeNER:
    """Mimics GLiNER's predict_entities well enough to prove the JD
    pipeline wiring (incl. the header/summary fix) actually reaches the
    model and comes back out through the label map."""

    def predict_entities(self, text, labels, threshold=0.3, flat_ner=True):
        entities = []
        if "Senior Data Engineer" in text:
            entities.append({"text": "Senior Data Engineer", "label": "job title"})
        if "Python" in text:
            entities.append({"text": "Python", "label": "programming language or technical skill"})
        if "SQL" in text:
            entities.append({"text": "SQL", "label": "programming language or technical skill"})
        return entities


class FakeEmbedder:
    def encode(self, texts, batch_size=1, max_length=512):
        import numpy as np
        # deterministic, distinct-ish vectors so cosine ordering is stable
        return {"dense_vecs": np.array([[0.1, 0.2, 0.3] for _ in texts])}


class FakeReranker:
    def compute_score(self, pairs, normalize=True, max_length=1024):
        # deterministic score per pair so ordering is stable/testable
        return [0.9 - (0.1 * i) for i in range(len(pairs))]


class FakeModelRegistry:
    @classmethod
    def ner(cls):
        return FakeNER()

    @classmethod
    def get_embedder(cls):
        return FakeEmbedder()

    @classmethod
    def get_reranker(cls):
        return FakeReranker()


# ---------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------
@pytest.fixture(autouse=True)
def patch_everything(monkeypatch):
    import pipeline.jd_processor as jd_processor
    import pipeline.ranker as ranker_mod
    import routers.jobs as jobs_router_mod
    import routers.candidates as candidates_router_mod
    import routers.matching as matching_router_mod
    import pipeline.storage as storage_mod

    # storage -> in-memory fake, patched at every import site that uses it
    monkeypatch.setattr(jobs_router_mod, "storage", fake_storage)
    monkeypatch.setattr(candidates_router_mod, "storage", fake_storage)
    monkeypatch.setattr(matching_router_mod, "storage", fake_storage)

    # ML models -> fakes
    monkeypatch.setattr(jd_processor, "ModelRegistry", FakeModelRegistry)
    monkeypatch.setattr(ranker_mod, "ModelRegistry", FakeModelRegistry)
    monkeypatch.setattr(jobs_router_mod, "ModelRegistry", FakeModelRegistry)
    monkeypatch.setattr(candidates_router_mod, "ModelRegistry", FakeModelRegistry)

    # Celery -> run synchronously in-process instead of dispatching to a
    # real broker/worker
    class FakeAsyncTask:
        def __init__(self, task_id, result):
            self.id = task_id
            self._result = result

    class FakeDelayResult:
        def __init__(self, task_id):
            self.id = task_id

    tasks_store = {}

    def fake_delay(**kwargs):
        from pipeline import ranker
        job = fake_storage.jobs[kwargs["job_id"]]
        candidate_ids = kwargs.get("candidate_ids") or list(fake_storage.candidates.keys())
        records = [fake_storage.candidates[cid] for cid in candidate_ids if cid in fake_storage.candidates]
        for r in records:
            r["semantic_score"] = 0.5
        ranked = ranker.rank_candidates(
            jd_text=job["jd_text"],
            jd_extracted=job["extracted"],
            candidate_records=records,
            top_n=kwargs.get("top_n", 20),
        )
        for r in ranked:
            r.pop("cv_text", None)
        task_id = uuid.uuid4().hex
        tasks_store[task_id] = {
            "signal": "MATCH_SUCCESS",
            "job_id": kwargs["job_id"],
            "shortlisted": len(records),
            "results": ranked,
        }
        return FakeDelayResult(task_id)

    class FakeCeleryResult:
        def __init__(self, task_id):
            self.task_id = task_id

        @property
        def state(self):
            return "SUCCESS" if self.task_id in tasks_store else "PENDING"

        @property
        def result(self):
            return tasks_store.get(self.task_id)

        @property
        def info(self):
            return tasks_store.get(self.task_id)

    class FakeCelery:
        def AsyncResult(self, task_id):
            return FakeCeleryResult(task_id)

    fake_run_match = type("FakeRunMatch", (), {"delay": staticmethod(fake_delay)})()

    monkeypatch.setattr(matching_router_mod, "run_match", fake_run_match)
    monkeypatch.setattr(matching_router_mod, "celery", FakeCelery())

    yield


@pytest.fixture
def client():
    from main import app
    return TestClient(app)


# ---------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------
def test_cors_headers_present(client):
    resp = client.options(
        "/api/v1/jobs/",
        headers={
            "Origin": "https://example.com",
            "Access-Control-Request-Method": "POST",
        },
    )
    assert resp.status_code in (200, 204)
    assert resp.headers.get("access-control-allow-origin") == "*"


def test_create_job_detects_title_via_header_fix(client):
    jd_text = (
        "Senior Data Engineer\n\n"
        "About the role\n"
        "We are hiring a Senior Data Engineer to build our pipelines.\n\n"
        "Requirements\n"
        "5+ years experience with Python and SQL\n"
    )
    resp = client.post("/api/v1/jobs/", json={"description": jd_text})
    assert resp.status_code == 201
    data = resp.json()
    assert data["signal"] == "JD_PROCESSED_SUCCESS"
    assert "job_id" in data
    # This is the exact bug from Fix 1: without header/summary in the NER
    # input, this list comes back empty.
    assert "senior data engineer" in data["extracted"]["required job title"]
    assert "python" in data["extracted"]["required hard skill"]


def test_full_pipeline_job_to_ranked_shortlist(client):
    # 1) create job
    jd_text = (
        "Senior Data Engineer\n\n"
        "About the role\n"
        "We are hiring a Senior Data Engineer.\n\n"
        "Requirements\nPython and SQL required.\n"
    )
    job_resp = client.post("/api/v1/jobs/", json={"description": jd_text})
    assert job_resp.status_code == 201
    job_id = job_resp.json()["job_id"]

    # 2) list jobs shows it
    list_resp = client.get("/api/v1/jobs/")
    assert list_resp.status_code == 200
    assert any(j["job_id"] == job_id for j in list_resp.json()["jobs"])

    # 3) upload a candidate CV as a real .docx (only pdf/docx are
    #    supported content types by preprocessing.dispatcher)
    from docx import Document

    cv_doc = Document()
    cv_doc.add_paragraph("Jane Smith")
    cv_doc.add_paragraph("jane@example.com")
    cv_doc.add_paragraph("+1 555 000 1111")
    cv_doc.add_paragraph("Experience")
    cv_doc.add_paragraph("2019 to Present at Acme Corp")
    cv_doc.add_paragraph("Education")
    cv_doc.add_paragraph("BSc Computer Science")
    cv_doc.add_paragraph("Skills")
    cv_doc.add_paragraph("Python, SQL, Docker")
    cv_doc.add_paragraph("Languages")
    cv_doc.add_paragraph("English")
    buf = io.BytesIO()
    cv_doc.save(buf)
    buf.seek(0)

    files = {
        "file": (
            "resume.docx",
            buf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    upload_resp = client.post("/api/v1/candidates/upload", files=files)
    assert upload_resp.status_code == 200, upload_resp.text
    upload_data = upload_resp.json()
    assert upload_data["signal"] == "CV_PREPROCESSED_SUCCESS"
    candidate_id = upload_data["file_id"]
    assert "python" in upload_data["parsed"]["skills"]

    # 4) dispatch matching (candidate_ids explicit -> uses our fake pool)
    match_resp = client.post(
        f"/api/v1/jobs/{job_id}/match",
        json={"candidate_ids": [candidate_id], "top_n": 5},
    )
    assert match_resp.status_code == 202
    task_id = match_resp.json()["task_id"]

    # 5) poll task result (our fake celery resolves synchronously)
    task_resp = client.get(f"/api/v1/jobs/tasks/{task_id}")
    assert task_resp.status_code == 200
    task_data = task_resp.json()
    assert task_data["status"] == "done"
    results = task_data["result"]["results"]
    assert len(results) == 1
    assert results[0]["cv_id"] == candidate_id
    assert "hard_match" in results[0]


def test_get_job_not_found_returns_404(client):
    resp = client.get("/api/v1/jobs/does-not-exist")
    assert resp.status_code == 404


def test_reject_too_short_job_description(client):
    resp = client.post("/api/v1/jobs/", json={"description": "too short"})
    assert resp.status_code == 400
    assert resp.json()["signal"] == "JD_TEXT_TOO_SHORT_OR_MISSING"
